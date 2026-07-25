#!/usr/bin/env python3
"""
Document format conversion to markdown for knowledge-wiki skill.
Supports: PDF, DOCX, XLSX, CSV, JSON, HTML
"""

import sys
import re
import json
import csv
from pathlib import Path
from dataclasses import dataclass
from typing import Optional, List, Dict, Any

# Import config loader
sys.path.insert(0, str(Path(__file__).parent))
from config_loader import get_config

CONFIG = get_config()
VAULT_ROOT = CONFIG.vault_root
INBOX_ROOT = CONFIG.inbox_root
ARCHIVE_ROOT = CONFIG.archive_root

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import marker
    MARKER_AVAILABLE = True
except ImportError:
    MARKER_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False


@dataclass
class ConversionResult:
    success: bool
    markdown: str
    metadata: Dict[str, Any]
    errors: List[str]
    warnings: List[str]


def convert_pdf(filepath: Path) -> ConversionResult:
    """Convert PDF to markdown using marker-pdf (preferred) or PyMuPDF."""
    errors = []
    warnings = []
    metadata = {"source_format": "pdf"}

    # Try marker-pdf first (better structure preservation)
    if MARKER_AVAILABLE:
        try:
            from marker.convert import convert_single_pdf
            markdown, _, images = convert_single_pdf(str(filepath))
            metadata["method"] = "marker-pdf"
            metadata["images_extracted"] = len(images) if images else 0
            return ConversionResult(True, markdown, metadata, [], warnings)
        except Exception as e:
            errors.append(f"marker-pdf failed: {e}")

    # Fallback to PyMuPDF
    if PYMUPDF_AVAILABLE:
        try:
            doc = fitz.open(str(filepath))
            text_parts = []
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    text_parts.append(f"## Page {page_num + 1}\n\n{text}")

            markdown = "\n\n".join(text_parts)
            metadata["method"] = "pymupdf"
            metadata["page_count"] = len(doc)
            doc.close()
            return ConversionResult(True, markdown, metadata, [], warnings)
        except Exception as e:
            errors.append(f"PyMuPDF failed: {e}")

    return ConversionResult(False, "", metadata, ["No PDF converter available (install marker-pdf or pymupdf)"], warnings)


def convert_docx(filepath: Path) -> ConversionResult:
    """Convert DOCX to markdown."""
    errors = []
    warnings = []
    metadata = {"source_format": "docx"}

    if not DOCX_AVAILABLE:
        return ConversionResult(False, "", metadata, ["python-docx not installed"], warnings)

    try:
        doc = Document(str(filepath))
        lines = []

        # Extract title from first heading or filename
        title_found = False
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                lines.append(f"{'#' * level} {para.text}")
                title_found = True
            elif para.text.strip():
                lines.append(para.text)
            else:
                lines.append("")

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            lines.append(f"\n### Table {table_idx + 1}\n")
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
                if row == table.rows[0]:
                    lines.append("| " + " | ".join(["---"] * len(cells)) + " |")

        markdown = "\n".join(lines)
        metadata["method"] = "python-docx"
        metadata["paragraph_count"] = len(doc.paragraphs)
        metadata["table_count"] = len(doc.tables)
        return ConversionResult(True, markdown, metadata, [], warnings)
    except Exception as e:
        return ConversionResult(False, "", metadata, [f"DOCX conversion failed: {e}"], warnings)


def convert_xlsx(filepath: Path) -> ConversionResult:
    """Convert XLSX/CSV to markdown tables."""
    errors = []
    warnings = []
    metadata = {"source_format": "xlsx"}

    if not OPENPYXL_AVAILABLE:
        return ConversionResult(False, "", metadata, ["openpyxl not installed"], warnings)

    try:
        wb = openpyxl.load_workbook(str(filepath), data_only=True)
        lines = [f"# {filepath.stem}\n"]

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            lines.append(f"\n## Sheet: {sheet_name}\n")

            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                lines.append("*Empty sheet*")
                continue

            # Convert to markdown table
            for row_idx, row in enumerate(rows):
                cells = [str(c) if c is not None else "" for c in row]
                lines.append("| " + " | ".join(cells) + " |")
                if row_idx == 0:
                    lines.append("| " + " | ".join(["---"] * len(cells)) + " |")

            lines.append("")

        markdown = "\n".join(lines)
        metadata["method"] = "openpyxl"
        metadata["sheet_count"] = len(wb.sheetnames)
        return ConversionResult(True, markdown, metadata, [], warnings)
    except Exception as e:
        return ConversionResult(False, "", metadata, [f"XLSX conversion failed: {e}"], warnings)


def convert_csv(filepath: Path) -> ConversionResult:
    """Convert CSV to markdown table."""
    errors = []
    warnings = []
    metadata = {"source_format": "csv"}

    try:
        with open(filepath, "r", encoding="utf-8") as f:
            reader = csv.reader(f)
            rows = list(reader)

        if not rows:
            return ConversionResult(True, "*Empty CSV*", metadata, [], warnings)

        lines = [f"# {filepath.stem}\n"]

        for row_idx, row in enumerate(rows):
            cells = [str(c) for c in row]
            lines.append("| " + " | ".join(cells) + " |")
            if row_idx == 0:
                lines.append("| " + " | ".join(["---"] * len(cells)) + " |")

        markdown = "\n".join(lines)
        metadata["method"] = "csv"
        metadata["row_count"] = len(rows)
        metadata["col_count"] = len(rows[0]) if rows else 0
        return ConversionResult(True, markdown, metadata, [], warnings)
    except Exception as e:
        return ConversionResult(False, "", metadata, [f"CSV conversion failed: {e}"], warnings)


def convert_json(filepath: Path) -> ConversionResult:
    """Convert JSON to markdown (formatted code block)."""
    errors = []
    warnings = []
    metadata = {"source_format": "json"}

    try:
        with open(filepath, "r", encoding="utf-8") as f:
            data = json.load(f)

        json_str = json.dumps(data, indent=2)
        markdown = f"# {filepath.stem}\n\n```json\n{json_str}\n```"

        metadata["method"] = "json"
        metadata["key_count"] = len(data) if isinstance(data, dict) else "N/A"
        return ConversionResult(True, markdown, metadata, [], warnings)
    except Exception as e:
        return ConversionResult(False, "", metadata, [f"JSON conversion failed: {e}"], warnings)


def convert_html(filepath: Path) -> ConversionResult:
    """Convert HTML to markdown (basic)."""
    errors = []
    warnings = []
    metadata = {"source_format": "html"}

    try:
        content = filepath.read_text(encoding="utf-8")
        # Very basic HTML -> markdown conversion
        text = re.sub(r"<h([1-6])[^>]*>(.*?)</h\1>", lambda m: f"{'#' * int(m.group(1))} {m.group(2)}", content, flags=re.DOTALL)
        text = re.sub(r"<p[^>]*>(.*?)</p>", r"\1\n\n", text, flags=re.DOTALL)
        text = re.sub(r"<br\s*/?>", "\n", text)
        text = re.sub(r"<strong[^>]*>(.*?)</strong>", r"**\1**", text, flags=re.DOTALL)
        text = re.sub(r"<em[^>]*>(.*?)</em>", r"*\1*", text, flags=re.DOTALL)
        text = re.sub(r"<code[^>]*>(.*?)</code>", r"`\1`", text, flags=re.DOTALL)
        text = re.sub(r"<pre[^>]*>(.*?)</pre>", r"```\n\1\n```", text, flags=re.DOTALL)
        text = re.sub(r"<a[^>]*href=[\"']([^\"']+)[\"'][^>]*>(.*?)</a>", r"[\2](\1)", text, flags=re.DOTALL)
        text = re.sub(r"<[^>]+>", "", text)  # Remove remaining tags

        markdown = f"# {filepath.stem}\n\n{text.strip()}"
        metadata["method"] = "regex"
        return ConversionResult(True, markdown, metadata, [], warnings)
    except Exception as e:
        return ConversionResult(False, "", metadata, [f"HTML conversion failed: {e}"], warnings)


def convert_to_markdown(filepath: Path) -> ConversionResult:
    """
    Main entry point - dispatches by file extension.
    Returns markdown content suitable for knowledge vault page.
    """
    ext = filepath.suffix.lower()

    converters = {
        ".pdf": convert_pdf,
        ".docx": convert_docx,
        ".doc": convert_docx,
        ".xlsx": convert_xlsx,
        ".xls": convert_xlsx,
        ".csv": convert_csv,
        ".json": convert_json,
        ".html": convert_html,
        ".htm": convert_html,
    }

    if ext not in converters:
        return ConversionResult(
            False, "",
            {"source_format": ext},
            [f"Unsupported format: {ext}"],
            []
        )

    result = converters[ext](filepath)

    # Add source file info to metadata
    result.metadata["source_file"] = str(filepath)
    result.metadata["source_filename"] = filepath.name

    return result


def batch_convert(input_dir: Path, output_dir: Path) -> List[ConversionResult]:
    """Convert all supported files in input_dir to markdown in output_dir."""
    results = []
    supported_exts = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".csv", ".json", ".html", ".htm"}

    for filepath in input_dir.rglob("*"):
        if filepath.is_file() and filepath.suffix.lower() in supported_exts:
            result = convert_to_markdown(filepath)
            if result.success:
                # Create output filename
                rel_path = filepath.relative_to(input_dir)
                out_name = rel_path.with_suffix(".md")
                out_path = output_dir / out_name
                out_path.parent.mkdir(parents=True, exist_ok=True)

                # Write with frontmatter
                frontmatter = {
                    "title": filepath.stem,
                    "created": filepath.stat().st_mtime,
                    "updated": filepath.stat().st_mtime,
                    "type": "entity",  # default, should be updated
                    "category": "unknown",  # must be set
                    "subcategory": "imported",
                    "tags": ["imported", filepath.suffix[1:]],
                    "sources": [{"title": filepath.name, "type": "file", "path": str(filepath)}],
                    "related": [],
                    "status": "draft",
                    "confidence": "low",
                    "version": "1.0",
                }

                import yaml
                fm_str = yaml.dump(frontmatter, default_flow_style=False)
                out_path.write_text(f"---\n{fm_str}---\n\n{result.markdown}", encoding="utf-8")

            results.append(result)

    return results


if __name__ == "__main__":
    import sys
    import yaml

    if len(sys.argv) < 2:
        print("Usage: python converter.py <file> [--batch <input_dir> <output_dir>]")
        sys.exit(1)

    if sys.argv[1] == "--batch":
        results = batch_convert(Path(sys.argv[2]), Path(sys.argv[3]))
        for r in results:
            print(yaml.dump({
                "success": r.success,
                "source": r.metadata.get("source_file"),
                "errors": r.errors,
                "warnings": r.warnings
            }))
    else:
        result = convert_to_markdown(Path(sys.argv[1]))
        print(yaml.dump({
            "success": result.success,
            "markdown": result.markdown[:500] + "..." if len(result.markdown) > 500 else result.markdown,
            "metadata": result.metadata,
            "errors": result.errors,
            "warnings": result.warnings
        }))